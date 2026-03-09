"""Tests for PDF/DOCX extraction, text normalization, and orchestrator.

Covers important (validation gate) criteria:
- PDF extraction produces correct page_num and source_unit_kind per page
- DOCX extraction covers paragraphs and table cells with sequential indices
- Norwegian characters (æøå) preserved after normalization
- Mojibake repaired (Ã¸ → ø)
- Control chars removed, newlines preserved, whitespace collapsed
- Empty/None input handled gracefully
- page_count backfilled in docs.parquet after extraction
- Warnings emitted for weak extraction but pipeline continues
"""

from pathlib import Path

import fitz
import pyarrow.parquet as pq
import pytest
from docx import Document

from src.ingestion.extraction import extract_docx_units, extract_pdf_units
from src.ingestion.normalization import normalize_text
from src.ingestion.run import (
    run_discovery_and_registration,
    run_extraction_and_normalization,
)


# ---------------------------------------------------------------------------
# Fixtures — synthetic PDF and DOCX files
# ---------------------------------------------------------------------------

@pytest.fixture()
def pdf_path(tmp_path: Path) -> Path:
    """Create a 2-page PDF with Norwegian text."""
    path = tmp_path / "test.pdf"
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((50, 72), "Side 1: Politiet etterforsker saken. Æøå.")
    p2 = doc.new_page()
    p2.insert_text((50, 72), "Side 2: Ola Nordmann er mistenkt.")
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture()
def empty_page_pdf(tmp_path: Path) -> Path:
    """Create a PDF where the second page is blank."""
    path = tmp_path / "empty_page.pdf"
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((50, 72), "Page with content.")
    doc.new_page()  # blank page
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture()
def docx_path(tmp_path: Path) -> Path:
    """Create a DOCX with paragraphs and Norwegian text."""
    path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Avhør av vitne. Æøå.")
    doc.add_paragraph("Vitnet forklarte at Kari Nordmann var til stede.")
    doc.save(str(path))
    return path


@pytest.fixture()
def docx_with_table(tmp_path: Path) -> Path:
    """Create a DOCX with paragraphs and a table."""
    path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    table.cell(1, 0).text = "Cell C"
    table.cell(1, 1).text = "Cell D"
    doc.save(str(path))
    return path


@pytest.fixture()
def case_with_real_files(tmp_path: Path) -> Path:
    """Case folder with a real PDF and a real DOCX for end-to-end tests."""
    root = tmp_path / "case"
    root.mkdir()

    # PDF
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((50, 72), "Rapport: DNB ASA er involvert.")
    pdf.save(str(root / "rapport.pdf"))
    pdf.close()

    # DOCX
    doc = Document()
    doc.add_paragraph("Avhør nummer 42.")
    doc.add_paragraph("Kari Nordmann forklarte seg.")
    doc.save(str(root / "avhoer.docx"))

    return root


@pytest.fixture()
def data_dir(tmp_path: Path) -> Path:
    out = tmp_path / "output"
    out.mkdir()
    return out


# ---------------------------------------------------------------------------
# PDF extraction tests
# ---------------------------------------------------------------------------

class TestExtractPdfUnits:
    def test_correct_page_count(self, pdf_path: Path):
        units = extract_pdf_units(pdf_path)
        assert len(units) == 2

    def test_page_nums_sequential(self, pdf_path: Path):
        units = extract_pdf_units(pdf_path)
        assert [u["page_num"] for u in units] == [0, 1]

    def test_source_unit_kind_is_pdf_page(self, pdf_path: Path):
        units = extract_pdf_units(pdf_path)
        assert all(u["source_unit_kind"] == "pdf_page" for u in units)

    def test_text_contains_norwegian_chars(self, pdf_path: Path):
        units = extract_pdf_units(pdf_path)
        combined = " ".join(u["text"] for u in units)
        assert "Æøå" in combined

    def test_text_is_non_empty(self, pdf_path: Path):
        units = extract_pdf_units(pdf_path)
        for u in units:
            assert u["text"].strip() != ""

    def test_empty_page_still_included(self, empty_page_pdf: Path):
        """Empty pages are kept as units (not dropped), just with empty text."""
        units = extract_pdf_units(empty_page_pdf)
        assert len(units) == 2
        assert units[1]["text"].strip() == ""

    def test_unit_keys(self, pdf_path: Path):
        units = extract_pdf_units(pdf_path)
        expected_keys = {"page_num", "text", "source_unit_kind"}
        for u in units:
            assert set(u.keys()) == expected_keys


# ---------------------------------------------------------------------------
# DOCX extraction tests
# ---------------------------------------------------------------------------

class TestExtractDocxUnits:
    def test_correct_paragraph_count(self, docx_path: Path):
        units = extract_docx_units(docx_path)
        assert len(units) == 2

    def test_page_nums_sequential(self, docx_path: Path):
        units = extract_docx_units(docx_path)
        assert [u["page_num"] for u in units] == [0, 1]

    def test_source_unit_kind_is_docx_paragraph(self, docx_path: Path):
        units = extract_docx_units(docx_path)
        assert all(u["source_unit_kind"] == "docx_paragraph" for u in units)

    def test_text_contains_norwegian_chars(self, docx_path: Path):
        units = extract_docx_units(docx_path)
        combined = " ".join(u["text"] for u in units)
        assert "Æøå" in combined

    def test_table_cells_appended_after_paragraphs(self, docx_with_table: Path):
        units = extract_docx_units(docx_with_table)
        # 1 paragraph + 4 table cells = 5 units
        assert len(units) == 5

    def test_table_cell_page_nums_continue_from_paragraphs(
        self, docx_with_table: Path
    ):
        units = extract_docx_units(docx_with_table)
        page_nums = [u["page_num"] for u in units]
        # paragraph at 0, then table cells at 1, 2, 3, 4
        assert page_nums == [0, 1, 2, 3, 4]

    def test_table_cell_text_present(self, docx_with_table: Path):
        units = extract_docx_units(docx_with_table)
        cell_texts = {u["text"] for u in units[1:]}
        assert "Cell A" in cell_texts
        assert "Cell D" in cell_texts

    def test_unit_keys(self, docx_path: Path):
        units = extract_docx_units(docx_path)
        expected_keys = {"page_num", "text", "source_unit_kind"}
        for u in units:
            assert set(u.keys()) == expected_keys


# ---------------------------------------------------------------------------
# Normalization tests
# ---------------------------------------------------------------------------

class TestNormalizeText:
    def test_preserves_norwegian_chars(self):
        assert "æøå" in normalize_text("æøå")
        assert "ÆØÅ" in normalize_text("ÆØÅ")

    def test_repairs_mojibake(self):
        # Common mojibake for ø in Latin-1 → UTF-8
        assert normalize_text("Ã¸") == "ø"

    def test_collapses_whitespace(self):
        assert normalize_text("hello   world") == "hello world"

    def test_strips_leading_trailing(self):
        assert normalize_text("  hello  ") == "hello"

    def test_removes_control_chars(self):
        assert normalize_text("hello\x00\x01world") == "helloworld"

    def test_preserves_newlines(self):
        result = normalize_text("hello\nworld")
        assert "\n" in result

    def test_empty_string(self):
        assert normalize_text("") == ""

    def test_none_input(self):
        assert normalize_text(None) == ""

    def test_tabs_converted_to_spaces(self):
        # Tabs become spaces to preserve word boundaries, then collapse
        assert normalize_text("hello\tworld") == "hello world"
        assert normalize_text("a\t\tb") == "a b"

    def test_nfc_normalization(self):
        import unicodedata
        # Decomposed ø (o + combining stroke) should become composed ø
        decomposed = unicodedata.normalize("NFD", "ø")
        result = normalize_text(decomposed)
        assert result == unicodedata.normalize("NFC", "ø")


# ---------------------------------------------------------------------------
# End-to-end orchestrator tests
# ---------------------------------------------------------------------------

class TestRunExtractionAndNormalization:
    def _setup_run(self, case_root: Path, data_dir: Path) -> str:
        """Helper: run discovery + registration and return run_id."""
        return run_discovery_and_registration(
            case_root, data_dir, run_id="e2e_extract"
        )

    def test_returns_dict_per_doc(self, case_with_real_files: Path, data_dir: Path):
        run_id = self._setup_run(case_with_real_files, data_dir)
        result = run_extraction_and_normalization(
            case_with_real_files, data_dir, run_id
        )
        assert len(result) == 2  # 1 PDF + 1 DOCX

    def test_units_have_required_keys(
        self, case_with_real_files: Path, data_dir: Path
    ):
        run_id = self._setup_run(case_with_real_files, data_dir)
        result = run_extraction_and_normalization(
            case_with_real_files, data_dir, run_id
        )
        expected_keys = {"page_num", "text", "source_unit_kind"}
        for units in result.values():
            for u in units:
                assert set(u.keys()) == expected_keys

    def test_text_is_normalized(self, case_with_real_files: Path, data_dir: Path):
        """Extracted text should not contain control chars (except newline)."""
        run_id = self._setup_run(case_with_real_files, data_dir)
        result = run_extraction_and_normalization(
            case_with_real_files, data_dir, run_id
        )
        for units in result.values():
            for u in units:
                # No control chars except \n
                for ch in u["text"]:
                    if ch != "\n":
                        assert ord(ch) >= 0x20 or ch == "\n"

    def test_page_count_backfilled(self, case_with_real_files: Path, data_dir: Path):
        run_id = self._setup_run(case_with_real_files, data_dir)
        run_extraction_and_normalization(case_with_real_files, data_dir, run_id)

        table = pq.read_table(data_dir / "docs.parquet")
        for row in table.to_pylist():
            assert row["page_count"] is not None
            assert row["page_count"] > 0

    def test_page_count_matches_unit_count(
        self, case_with_real_files: Path, data_dir: Path
    ):
        run_id = self._setup_run(case_with_real_files, data_dir)
        result = run_extraction_and_normalization(
            case_with_real_files, data_dir, run_id
        )

        table = pq.read_table(data_dir / "docs.parquet")
        for row in table.to_pylist():
            doc_id = row["doc_id"]
            assert row["page_count"] == len(result[doc_id])

    def test_docs_parquet_still_contract_valid(
        self, case_with_real_files: Path, data_dir: Path
    ):
        """docs.parquet remains schema-valid after page_count backfill."""
        from src.shared.schemas import validate_contract_rules

        run_id = self._setup_run(case_with_real_files, data_dir)
        run_extraction_and_normalization(case_with_real_files, data_dir, run_id)

        table = pq.read_table(data_dir / "docs.parquet")
        errors = validate_contract_rules(table, "docs")
        assert errors == []

    def test_missing_docs_parquet_returns_empty(self, tmp_path: Path):
        result = run_extraction_and_normalization(
            tmp_path, tmp_path / "nonexistent", "norun"
        )
        assert result == {}

    def test_unknown_run_id_returns_empty(
        self, case_with_real_files: Path, data_dir: Path
    ):
        self._setup_run(case_with_real_files, data_dir)
        result = run_extraction_and_normalization(
            case_with_real_files, data_dir, "nonexistent_run"
        )
        assert result == {}

    def test_changed_file_skips_extraction_preserves_provenance(
        self, tmp_path: Path, data_dir: Path
    ):
        """Rerun after in-place file change must not corrupt the docs row.

        Scenario from code review: ingest a 1-page PDF with run_id=R,
        replace the file with a 3-page PDF, rerun with run_id=R.
        The docs row must keep page_count=1 (from the original extraction),
        and the changed file must be skipped entirely.
        """
        case_root = tmp_path / "case"
        case_root.mkdir()
        run_id = "same_run"

        # Create a 1-page PDF and run full discovery+registration+extraction
        pdf_path = case_root / "report.pdf"
        doc1 = fitz.open()
        p = doc1.new_page()
        p.insert_text((50, 72), "Original single page.")
        doc1.save(str(pdf_path))
        doc1.close()

        run_discovery_and_registration(case_root, data_dir, run_id=run_id)
        result1 = run_extraction_and_normalization(
            case_root, data_dir, run_id
        )
        assert len(result1) == 1
        original_doc_id = list(result1.keys())[0]

        # Verify page_count is 1 after first extraction
        table1 = pq.read_table(data_dir / "docs.parquet")
        row1 = table1.to_pylist()[0]
        assert row1["page_count"] == 1
        assert row1["doc_id"] == original_doc_id

        # Replace the file with a 3-page PDF (different content)
        doc2 = fitz.open()
        for i in range(3):
            pg = doc2.new_page()
            pg.insert_text((50, 72), f"Replacement page {i + 1}.")
        doc2.save(str(pdf_path))
        doc2.close()

        # Rerun extraction with same run_id — changed file must be skipped
        result2 = run_extraction_and_normalization(
            case_root, data_dir, run_id
        )
        assert original_doc_id not in result2

        # docs.parquet must still show the original page_count
        table2 = pq.read_table(data_dir / "docs.parquet")
        row2 = table2.to_pylist()[0]
        assert row2["doc_id"] == original_doc_id
        assert row2["page_count"] == 1  # NOT 3
