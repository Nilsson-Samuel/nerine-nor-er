"""Text extraction — PDF via PyMuPDF and DOCX via python-docx.

Returns position-aware text units (page/paragraph-level) for downstream
normalization and chunking. Warns on empty/weak extraction but never crashes.
"""

import logging
from pathlib import Path

import fitz  # PyMuPDF

from docx import Document

logger = logging.getLogger(__name__)


def extract_pdf_units(path: Path) -> list[dict]:
    """Extract text units from a PDF, one per page.

    Args:
        path: Absolute path to the PDF file.

    Returns:
        List of dicts with keys: page_num, text, source_unit_kind.
        Empty pages produce units with empty text (warned, not skipped).
    """
    units = []
    doc = fitz.open(path)
    for page_num, page in enumerate(doc):
        text = page.get_text() or ""
        if not text.strip():
            logger.warning("empty_pdf_page: %s page %d", path.name, page_num)
        units.append({
            "page_num": page_num,
            "text": text,
            "source_unit_kind": "pdf_page",
        })
    doc.close()
    return units


def extract_docx_units(path: Path) -> list[dict]:
    """Extract text units from a DOCX — paragraphs followed by table cells.

    Paragraphs use their natural index as page_num. Table cells are appended
    after paragraphs in deterministic row-major order.

    Args:
        path: Absolute path to the DOCX file.

    Returns:
        List of dicts with keys: page_num, text, source_unit_kind.
    """
    doc = Document(path)
    units = []

    # Paragraphs
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text or ""
        if not text.strip():
            logger.warning("empty_docx_paragraph: %s paragraph %d", path.name, idx)
        units.append({
            "page_num": idx,
            "text": text,
            "source_unit_kind": "docx_paragraph",
        })

    # Table cells — appended after paragraphs in row-major order
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                units.append({
                    "page_num": len(units),
                    "text": text,
                    "source_unit_kind": "docx_paragraph",
                })

    return units
